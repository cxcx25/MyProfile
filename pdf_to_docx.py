import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def pdf_to_docx(pdf_path, docx_path):
    # Create Word document
    doc = Document()
    
    # Extract images using PyMuPDF
    pdf_document = fitz.open(pdf_path)
    
    # Extract text and images from PDF with better formatting
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract images from the page using PyMuPDF
            page_pix = pdf_document[page_num]
            image_list = page_pix.get_images()
            
            if image_list:
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save image temporarily
                    image_path = f"temp_image_{page_num}_{img_index}.{image_ext}"
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    # Add image to document
                    try:
                        doc.add_picture(image_path, width=Inches(1.5))
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        pass
                    # Clean up temp image
                    if os.path.exists(image_path):
                        os.remove(image_path)
            
            text = page.extract_text()
            if text:
                # Split by lines and add as paragraphs
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:  # Skip empty lines
                        para = doc.add_paragraph(line)
                        # Try to detect headings (shorter lines, often in caps or bold-like)
                        if len(line) < 50 and line.isupper():
                            para.runs[0].bold = True
                            para.runs[0].font.size = Pt(14)
                # Add page break between pages (except last page)
                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()
    
    pdf_document.close()
    # Save the document
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    pdf_path = r"c:\Project\MyProfile\Joseph_Rey_Marilla_CV.pdf"
    docx_path = r"c:\Project\MyProfile\Joseph Rey Marilla.docx"
    pdf_to_docx(pdf_path, docx_path)
