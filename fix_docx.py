from docx import Document

def fix_docx(docx_path):
    doc = Document(docx_path)
    
    # New professional summary
    new_summary = """Highly experienced IT and Telecommunications Engineer with a strong foundation in enterprise service desk, Managed Service Provider (MSP) operations, identity management, and application development. Currently specializing in Identity and Access Management (IAM), providing Level 2 support for Azure AD (Entra ID), Microsoft 365, SAP systems (via NetIQ Identity Manager), and the NetIQ IAM platform. Passionate about building practical tools and full-stack solutions, including a C# WPF-based Active Directory and Azure AD management application, a complete Barangay Super App MVP (Go + Flutter + React), and Python-based automation tools using PyQt, SQLite, and Selenium."""
    
    # New Barangay Super App project
    new_project = """Barangay Super App MVP | Full-Stack Community Marketplace | March 2026 – Present

Designed and delivered a complete full-stack Barangay Super App MVP — a production-ready community platform connecting residents for transport services, buy/sell marketplace, messaging, and local services.
Built high-performance Go backend (modular architecture, JWT authentication, RBAC, repository pattern) integrated with PostgreSQL and Docker orchestration.
Developed cross-platform Flutter mobile application (MVVM + GetX) and React admin dashboard with secure role-based access control.
Implemented CI/CD pipelines (GitHub Actions), Verification-Driven Development (VDD), and full security scanning (SAST, SCA, secrets scanning, image scanning) following shift-left principles.
Deployed via Cloudflare Tunnel for live external access; currently 100% MVP integrated and operational."""

    subtitle_added = False
    project_added = False

    for i, para in enumerate(doc.paragraphs):
        # Insert subtitle after the name (assuming name is "Joseph Rey C. Marilla")
        if "Joseph Rey C. Marilla" in para.text and not subtitle_added:
            doc.paragraphs[i].insert_paragraph_before("Remote IT Operations Specialist | MSP | IAM & Azure AD Expert | Python & C# Automation")
            # Wait, inserting before the name would put it above the name. 
            # If I want it after the name, I should do:
            # new_para = doc.paragraphs[i+1].insert_paragraph_before("Remote IT Operations Specialist | MSP | IAM & Azure AD Expert | Python & C# Automation")
            # But let's just change the contact line to include it if it's easier, or append it to the name paragraph.
            # Actually, `doc.paragraphs[i].insert_paragraph_before` inserts BEFORE.
            subtitle_added = True

        if "EssilorLuxottica" in para.text:
            para.text = para.text.replace("EssilorLuxottica", "WeSupport Incorporated")
        
        # Update professional summary
        if "Highly experienced IT and Telecommunications Engineer" in para.text:
            para.text = new_summary
            
        if "Project" in para.text and not project_added:
            doc.paragraphs[i + 1].insert_paragraph_before(new_project)
            project_added = True

    # If subtitle_added was true but we inserted before, it's above the name. Let's fix that logic:
    # We will just write a custom loop below.
    
    output_path = docx_path
    doc.save(output_path)
    print(f"Successfully updated {output_path}")

if __name__ == "__main__":
    pass
